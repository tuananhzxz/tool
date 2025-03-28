import time
from flask import Flask, render_template, request, jsonify, send_file, url_for
import os
from werkzeug.utils import secure_filename
from PIL import Image, ImageEnhance
import shutil
from changeImage import convert_images
from cutmergeimage import ImageProcessor
from mergeWord import merge_word_documents
from renameImage import rename_files
import requests
from bs4 import BeautifulSoup
import io
import zipfile
import re
from flask_cors import CORS
from ocr_processor import process_folder
from add_logo import LogoProcessor
from docx import Document
from google import genai
from google.genai import types
from dotenv import load_dotenv
from service.dowloadImg import (
    download_selected_images,
)

# Load environment variables
load_dotenv()

app = Flask(__name__, static_folder='public')
CORS(app)
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB max-limit

# Hàm xử lý lỗi chung
def handle_error(e):
    error_msg = str(e)
    if isinstance(e, requests.exceptions.RequestException):
        error_msg = "Lỗi kết nối: Không thể tải dữ liệu từ URL"
    elif isinstance(e, IOError):
        error_msg = "Lỗi đọc/ghi file"
    elif isinstance(e, Image.UnidentifiedImageError):
        error_msg = "File ảnh không hợp lệ hoặc bị hỏng"
    return jsonify({'error': error_msg})

@app.route('/')
def index():
    functions = {
        'changeImage': 'Chuyển đổi định dạng ảnh',
        'cutmergeimage': 'Cắt và ghép ảnh',
        'dowloaf': 'Tải ảnh từ web',
        'mergeWord': 'Gộp file Word',
        'renameImage': 'Đổi tên ảnh',
        'ocr': 'Nhận dạng chữ trong ảnh (OCR)',
        'addlogo': 'Thêm logo vào ảnh'
    }
    return render_template('index.html', functions=functions)

@app.route('/function/<function_name>')
def function_page(function_name):
    templates = {
        'changeImage': 'functions/changeimage.html',
        'cutmergeimage': 'functions/cutmergeimage.html',
        'dowloaf': 'functions/dowloaf.html',
        'mergeWord': 'functions/mergeword.html',
        'renameImage': 'functions/renameimage.html',
        'ocr': 'functions/ocr.html',
        'translate': 'functions/translate.html',
        'addlogo': 'functions/addlogo.html'  # Add this line
    }
    
    if function_name not in templates:
        return jsonify({'error': 'Chức năng không tồn tại'}), 404
        
    return render_template(templates[function_name])

@app.route('/execute/changeImage', methods=['POST'])
def execute_change_image():
    try:
        if 'files' not in request.files:
            return jsonify({'error': 'Không có file được tải lên'})
        
        files = request.files.getlist('files')
        target_format = request.form.get('target_format', 'webp')
        
        if not files or files[0].filename == '':
            return jsonify({'error': 'Không có file được chọn'})
        
        # Tạo file zip trong bộ nhớ
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w') as zipf:
            # Theo dõi các chuyển đổi đã thực hiện
            conversions = []
            
            # Xử lý từng file
            for file in files:
                if file.filename:
                    # Đọc file vào bộ nhớ
                    image_bytes = file.read()
                    img = Image.open(io.BytesIO(image_bytes))
                    
                    # Xác định định dạng nguồn
                    source_format = img.format.lower() if img.format else os.path.splitext(file.filename)[1][1:].lower()
                    conversions.append(f"{source_format.upper()} → {target_format.upper()}")
                    
                    # Chuyển đổi và lưu vào bộ nhớ
                    output_buffer = io.BytesIO()
                    if target_format.upper() == 'WEBP':
                        img.save(output_buffer, 'WEBP', quality=90)
                    elif target_format.upper() == 'JPEG':
                        if img.mode in ('RGBA', 'LA'):
                            background = Image.new('RGB', img.size, (255, 255, 255))
                            background.paste(img, mask=img.split()[-1])
                            background.save(output_buffer, 'JPEG', quality=95)
                        else:
                            img.save(output_buffer, 'JPEG', quality=95)
                    else:  # PNG
                        img.save(output_buffer, 'PNG')
                    
                    # Thêm vào zip
                    output_buffer.seek(0)
                    zipf.writestr(
                        os.path.splitext(file.filename)[0] + '.' + target_format.lower(),
                        output_buffer.getvalue()
                    )
        
        # Tạo thông báo chi tiết về các chuyển đổi
        unique_conversions = list(set(conversions))
        conversion_message = ", ".join(unique_conversions)
        
        # Gửi file zip
        zip_buffer.seek(0)
        return send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name='converted_images.zip'
        )
        
    except Exception as e:
        return handle_error(e)

@app.route('/execute/cutmergeimage', methods=['POST'])
def execute_cut_merge_image():
    if 'files' not in request.files:
        return jsonify({'error': 'Không có file được tải lên'})
    
    files = request.files.getlist('files')
    action = request.form.get('action')
    min_height = int(request.form.get('min_height', 2500))
    images_per_group = int(request.form.get('parts', 2))
    width = request.form.get('width')
    height = request.form.get('height')
    
    if not files or files[0].filename == '':
        return jsonify({'error': 'Không có file được chọn'})
    
    try:
        # Xử lý ảnh trong bộ nhớ
        processed_images = []
        for file in files:
            if file.filename:
                # Đọc ảnh vào bộ nhớ
                image_bytes = file.read()
                img = Image.open(io.BytesIO(image_bytes))
                
                # Resize nếu cần
                if width or height:
                    new_width = int(width) if width else img.width
                    new_height = int(height) if height else img.height
                    img = img.resize((new_width, new_height))
                
                processed_images.append(img)
        
        # Xử lý ảnh
        if action == 'split':
            # Tạo file zip trong bộ nhớ
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w') as zipf:
                for idx, img in enumerate(processed_images):
                    # Cắt ảnh
                    height = img.height
                    parts = height // min_height
                    for i in range(parts):
                        top = i * min_height
                        bottom = (i + 1) * min_height
                        part = img.crop((0, top, img.width, bottom))
                        
                        # Lưu phần cắt vào bộ nhớ
                        output_buffer = io.BytesIO()
                        part.save(output_buffer, format='PNG')
                        output_buffer.seek(0)
                        
                        # Thêm vào zip
                        zipf.writestr(f"{idx + 1}_{i + 1}.png", output_buffer.getvalue())
        else:  # merge
            if len(processed_images) < images_per_group:
                return jsonify({'error': f'Số lượng ảnh ({len(processed_images)}) phải lớn hơn hoặc bằng số ảnh mỗi nhóm ({images_per_group})'})
            
            # Tạo file zip trong bộ nhớ
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w') as zipf:
                # Ghép ảnh theo nhóm
                for i in range(0, len(processed_images), images_per_group):
                    group = processed_images[i:i + images_per_group]
                    # Tính kích thước cho ảnh ghép
                    total_height = sum(img.height for img in group)
                    max_width = max(img.width for img in group)
                    
                    # Tạo ảnh mới
                    merged = Image.new('RGB', (max_width, total_height))
                    
                    # Ghép ảnh
                    y_offset = 0
                    for img in group:
                        merged.paste(img, (0, y_offset))
                        y_offset += img.height
                    
                    # Lưu ảnh ghép vào bộ nhớ
                    output_buffer = io.BytesIO()
                    merged.save(output_buffer, format='PNG')
                    output_buffer.seek(0)
                    
                    # Thêm vào zip
                    zipf.writestr(f"merged_{i//images_per_group + 1}.png", output_buffer.getvalue())
        
        # Gửi file zip
        zip_buffer.seek(0)
        return send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name='processed_images.zip'
        )
        
    except Exception as e:
        return handle_error(e)

@app.route('/execute/mergeWord', methods=['POST'])
def execute_merge_word():
    if 'files' not in request.files:
        return jsonify({'error': 'Không có file được tải lên'})
    
    files = request.files.getlist('files')
    if not files or files[0].filename == '':
        return jsonify({'error': 'Không có file được chọn'})
    
    try:
        # Tạo document mới
        merged_doc = Document()
        
        # Gộp các file Word
        for file in files:
            if file.filename:
                doc = Document(file)
                for element in doc.element.body:
                    merged_doc.element.body.append(element)
        
        # Lưu vào bộ nhớ
        output_buffer = io.BytesIO()
        merged_doc.save(output_buffer)
        output_buffer.seek(0)
        
        return send_file(
            output_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='merged_document.docx'
        )
        
    except Exception as e:
        return handle_error(e)

@app.route('/execute/renameImage', methods=['POST'])
def execute_rename_image():
    if 'files' not in request.files:
        return jsonify({'error': 'Không có file được tải lên'})
    
    files = request.files.getlist('files')
    if not files or files[0].filename == '':
        return jsonify({'error': 'Không có file được chọn'})
    
    try:
        # Tạo file zip trong bộ nhớ
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w') as zipf:
            # Xử lý từng file
            for idx, file in enumerate(files, 1):
                if file.filename:
                    # Đọc file vào bộ nhớ
                    file_bytes = file.read()
                    
                    # Tạo tên mới
                    ext = os.path.splitext(file.filename)[1]
                    new_name = f"{idx}{ext}"
                    
                    # Thêm vào zip
                    zipf.writestr(new_name, file_bytes)
        
        # Gửi file zip
        zip_buffer.seek(0)
        return send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name='renamed_images.zip'
        )
        
    except Exception as e:
        return handle_error(e)

@app.route('/execute/ocr', methods=['POST'])
def execute_ocr():
    try:
        files = request.files.getlist('files')
        api_key = request.form.get('api_key')
        mode = request.form.get('mode', 'ocr')
        genres = request.form.getlist('genres[]')
        styles = request.form.getlist('styles[]')
        
        if not files or not api_key:
            return jsonify({'error': 'Vui lòng cung cấp file và API key'})
        
        # Tạo một tài liệu Word mới
        doc = Document()
        
        processed_files = []
        failed_files = []
        
        try:
            # Xử lý từng file một
            for file in files:
                if not file.filename:
                    continue
                    
                try:
                    # Kiểm tra định dạng file
                    file_ext = os.path.splitext(file.filename)[1].lower()
                    
                    if file_ext in ['.docx', '.doc']:
                        # Xử lý file Word
                        success = process_word_file(file, doc, mode, genres, styles, api_key)
                        if success:
                            processed_files.append(file.filename)
                        else:
                            failed_files.append(file.filename)
                            
                    elif file_ext in ['.jpg', '.jpeg', '.png', '.webp']:
                        # Xử lý file ảnh trực tiếp từ bộ nhớ
                        success = process_image_file_memory(file, doc, mode, genres, styles, api_key)
                        if success:
                            processed_files.append(file.filename)
                        else:
                            failed_files.append(file.filename)
                    else:
                        failed_files.append(file.filename)
                        continue
                        
                except Exception as e:
                    print(f"Lỗi khi xử lý file {file.filename}: {str(e)}")
                    failed_files.append(file.filename)
                    continue
            
            if not processed_files:
                return jsonify({'error': 'Không thể xử lý bất kỳ file nào'})
            
            # Lưu tài liệu vào bộ nhớ
            output_filename = f'VBCĐ_{int(time.time())}.docx'
            output_buffer = io.BytesIO()
            doc.save(output_buffer)
            output_buffer.seek(0)
            
            return send_file(
                output_buffer,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=output_filename
            )
            
        except Exception as e:
            return jsonify({'error': str(e)})
        
    except Exception as e:
        return jsonify({'error': str(e)})

def process_word_file(file, doc, mode, genres, styles, api_key):
    try:
        docx_doc = Document(file)
        text = ''
        for para in docx_doc.paragraphs:
            text += para.text + '\n'
        
        if text.strip():
            if mode == 'translate':
                # Dịch văn bản
                target_langs = request.form.getlist('target_langs[]')
                for lang in target_langs:
                    prompt = f"""
                        Hãy dịch đoạn văn bản sau sang {getLangName(lang)}.
                        Thể loại: {', '.join(genres) if genres else 'Không xác định'}
                        Phong cách: {', '.join(styles) if styles else 'Không xác định'}
                        
                        Văn bản cần dịch:
                        {text}
                    """
                    client = genai.Client(api_key=api_key)
                    
                    response = client.models.generate_content(
                        model="gemini-2.5-pro-exp-03-25",
                        contents=prompt,
                        config={"temperature": 0.2}
                    )
                    
                    if response.text:
                        doc.add_paragraph(f"=== {getLangName(lang)} ===")
                        doc.add_paragraph(response.text)
                        doc.add_paragraph('---')
            else:
                doc.add_paragraph(text)
                doc.add_paragraph('---')
        return True
    except Exception as e:
        print(f"Lỗi khi xử lý file Word {file.filename}: {str(e)}")
        return False

def process_image_file_memory(file, doc, mode, genres, styles, api_key):
    try:
        # Đọc ảnh từ bộ nhớ
        image_bytes = file.read()
        img = Image.open(io.BytesIO(image_bytes))
        
        # Kiểm tra kích thước ảnh
        max_size = (1920, 1080)
        if img.size[0] > max_size[0] or img.size[1] > max_size[1]:
            img.thumbnail(max_size, Image.Resampling.LANCZOS)
        
        # Chuyển sang grayscale và tăng độ tương phản
        img = img.convert('L')
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(1.5)
        
        # Chuyển ảnh đã xử lý thành bytes
        img_byte_arr = io.BytesIO()
        img.save(img_byte_arr, format='PNG')
        img_byte_arr = img_byte_arr.getvalue()
        
        # Tạo prompt cho việc nhận dạng văn bản
        prompt = f"""
            NHIỆM VỤ: Nhận dạng và trích xuất văn bản từ ảnh với độ chính xác cao nhất.
            
            YÊU CẦU CHẤT LƯỢNG:
            1. Nhận dạng chính xác 100% nội dung văn bản, kể cả chữ nhỏ
            2. Phân biệt rõ các đoạn văn bản khác nhau, các bóng thoại khác nhau
            3. Giữ nguyên vị trí và thứ tự của các bóng thoại
            4. Không bỏ sót bất kỳ ký tự nào
            
            QUY TẮC XỬ LÝ:
            1. Loại bỏ các yếu tố không phải văn bản 
            2. QUAN TRỌNG: Xử lý mỗi bóng thoại (speech bubble) như MỘT CÂU HOÀN CHỈNH TRÊN MỘT DÒNG DUY NHẤT
            3. Mỗi bóng thoại riêng biệt sẽ được xuất ra thành một dòng văn bản riêng biệt
            4. Giữ nguyên các dấu câu và định dạng đặc biệt
            
            ĐỊNH DẠNG ĐẦU RA:
            - Mỗi bóng thoại trên một dòng riêng
            - Giữ nguyên các dấu câu và định dạng
            - Không thêm bất kỳ chú thích hay giải thích nào
            - Không tách văn bản trong một bóng thoại thành nhiều dòng
        """
        
        client = genai.Client(api_key=api_key)
        
        # Gửi yêu cầu đến Gemini API với cấu hình tối ưu
        response = client.models.generate_content(
            model="gemini-2.5-pro-exp-03-25",
            contents=[
                {"text": prompt},
                {"inline_data": {"mime_type": "image/png", "data": img_byte_arr}}
            ],
            config={
                "temperature": 0.1,  # Giảm temperature để tăng độ chính xác
                "max_output_tokens": 2048,  # Tăng max tokens để xử lý văn bản dài
                "top_p": 0.8,
                "top_k": 40,
                "candidate_count": 1
            }
        )
        
        if response.text:
            # Thêm thông tin về file đang xử lý
            doc.add_paragraph(f"=== File: {file.filename} ===")
            
            if mode == 'translate':
                # Dịch văn bản
                target_langs = request.form.getlist('target_langs[]')
                for lang in target_langs:
                    prompt = f"""
                        Hãy dịch đoạn văn bản sau sang {getLangName(lang)}.
                        Thể loại: {', '.join(genres) if genres else 'Không xác định'}
                        Phong cách: {', '.join(styles) if styles else 'Không xác định'}
                        
                        Văn bản cần dịch:
                        {response.text}
                    """
                    
                    response = client.models.generate_content(
                        model="gemini-2.5-pro-exp-03-25",
                        contents=prompt,
                        config={
                            "temperature": 0.2,
                            "max_output_tokens": 2048
                        }
                    )
                    
                    if response.text:
                        doc.add_paragraph(f"=== {getLangName(lang)} ===")
                        doc.add_paragraph(response.text)
                        doc.add_paragraph('---')
            else:
                doc.add_paragraph(response.text)
                doc.add_paragraph('---')
                
            # Thêm delay giữa các lần xử lý để tránh rate limit
            time.sleep(1)
            
        return True
    except Exception as e:
        print(f"Lỗi khi xử lý file ảnh {file.filename}: {str(e)}")
        return False

# Hàm chuyển đổi mã ngôn ngữ thành tên
def getLangName(code):
    langs = {
        'vie': 'Tiếng Việt',
        'eng': 'Tiếng Anh',
        'jpn': 'Tiếng Nhật',
        'kor': 'Tiếng Hàn'
    }
    return langs.get(code, code)

@app.route('/execute/merge_ocr', methods=['POST'])
def execute_merge_ocr():
    if 'files' not in request.files:
        return jsonify({'error': 'Không có file được tải lên'})
    
    files = request.files.getlist('files')
    if not files or files[0].filename == '':
        return jsonify({'error': 'Không có file được chọn'})
    
    try:
        # Tạo document mới
        merged_doc = Document()
        
        # Gộp các file Word
        for file in files:
            if file.filename:
                doc = Document(file)
                for element in doc.element.body:
                    merged_doc.element.body.append(element)
        
        # Lưu vào bộ nhớ
        output_buffer = io.BytesIO()
        merged_doc.save(output_buffer)
        output_buffer.seek(0)
        
        return send_file(
            output_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='merged_ocr.docx'
        )
        
    except Exception as e:
        return handle_error(e)

@app.route('/download/<filename>')
def download_file(filename):
    try:
        file_path = os.path.join(os.getcwd(), filename)
        if not os.path.exists(file_path):
            return jsonify({'error': 'File không tồn tại'})
        return send_file(file_path, as_attachment=True)
    except Exception as e:
        return handle_error(e)

@app.errorhandler(413)
def too_large(e):
    return jsonify({'error': 'File quá lớn. Kích thước tối đa là 16MB'}), 413

@app.errorhandler(500)
def internal_error(e):
    return jsonify({'error': 'Lỗi máy chủ nội bộ'}), 500

@app.errorhandler(404)
def not_found(e):
    return jsonify({'error': 'Không tìm thấy trang'}), 404

@app.route('/changeImage')
def change_image():
    return render_template('functions/changeimage.html')

@app.route('/cutmergeimage')
def cut_merge_image():
    return render_template('functions/cutmergeimage.html')

@app.route('/dowloaf')
def download_image():
    return render_template('functions/dowloaf.html')

@app.route('/mergeword')
def merge_word():
    return render_template('functions/mergeword.html')

@app.route('/renameimage')
def rename_image():
    return render_template('functions/renameimage.html')

@app.route('/ocr')
def ocr():
    return render_template('functions/ocr.html')

@app.route('/translate')
def translate():
    return render_template('functions/translate.html')

@app.route('/remove_speech_bubbles')
def remove_speech_bubbles():
    return render_template('functions/speechbubble.html')

@app.route('/test_api_key', methods=['POST'])
def test_api_key():
    try:
        data = request.json
        api_key = data.get('api_key')
        
        if not api_key:
            return jsonify({'error': 'API key không được cung cấp'})
        
        client = genai.Client(api_key=api_key)
        
        # Tạo một tin nhắn đơn giản để kiểm tra API key
        response = client.models.generate_content(
            model="gemini-2.5-pro-exp-03-25", contents="Explain how AI works in a few words"
        )
        
        return jsonify({'success': True, 'message': 'API key hợp lệ'})
        
    except Exception as e:
        return jsonify({'error': str(e)})

@app.route('/execute/download_images', methods=['POST'])
def execute_download():
    return download_selected_images()

@app.route('/add_logo', methods=['GET', 'POST'])
def add_logo():
    if request.method == 'GET':
        return render_template('functions/addlogo.html')
    
    try:
        # Kiểm tra file upload
        if 'files' not in request.files or 'logo' not in request.files:
            return jsonify({'error': 'Vui lòng chọn cả logo và ảnh cần xử lý'}), 400
            
        files = request.files.getlist('files')
        logo = request.files['logo']
        
        if not files or not logo or files[0].filename == '' or logo.filename == '':
            return jsonify({'error': 'Vui lòng chọn file'}), 400
            
        # Lấy các tham số
        position = request.form.get('position', 'top_left')
        scale = float(request.form.get('scale', 10)) / 100
        
        try:
            # Đọc logo vào bộ nhớ
            logo_bytes = logo.read()
            logo_img = Image.open(io.BytesIO(logo_bytes))
            
            # Tính toán kích thước logo
            logo_width = int(logo_img.width * scale)
            logo_height = int(logo_img.height * scale)
            logo_img = logo_img.resize((logo_width, logo_height))
            
            # Tạo file zip trong bộ nhớ
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w') as zipf:
                # Xử lý từng ảnh
                for file in files:
                    if file.filename:
                        # Đọc ảnh vào bộ nhớ
                        image_bytes = file.read()
                        img = Image.open(io.BytesIO(image_bytes))
                        
                        # Tính vị trí logo
                        if position == 'top_left':
                            pos = (10, 10)
                        elif position == 'top_right':
                            pos = (img.width - logo_width - 10, 10)
                        elif position == 'bottom_left':
                            pos = (10, img.height - logo_height - 10)
                        else:  # bottom_right
                            pos = (img.width - logo_width - 10, img.height - logo_height - 10)
                        
                        # Thêm logo vào ảnh
                        img.paste(logo_img, pos, logo_img)
                        
                        # Lưu ảnh vào bộ nhớ
                        output_buffer = io.BytesIO()
                        img.save(output_buffer, format='PNG')
                        output_buffer.seek(0)
                        
                        # Thêm vào zip
                        zipf.writestr(file.filename, output_buffer.getvalue())
            
            # Gửi file zip
            zip_buffer.seek(0)
            return send_file(
                zip_buffer,
                mimetype='application/zip',
                as_attachment=True,
                download_name='processed_images.zip'
            )
            
        except Exception as e:
            return jsonify({'error': str(e)}), 500
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/addlogo')
def add_logo_page():
    return render_template('functions/addlogo.html')

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8000, debug=True)
